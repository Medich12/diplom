import datetime
import os
import zipfile
import django.http
import docx
from urllib.parse import quote
from django.core.files import File
import urllib.parse
from django.db.models import Max
from django.db.models.functions import TruncDate
from django.template.loader import render_to_string
from django.http import JsonResponse
from django.db.models import Q
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# import reportlab.lib.pagesizes
from bs4 import BeautifulSoup
import requests
from django.http import HttpResponse
from django.http import FileResponse
from django.views.generic import TemplateView, ListView
from django.shortcuts import render, redirect, get_object_or_404
from openpyxl.reader.excel import load_workbook

from kursachDjango import settings
from .models import Car, PhotoCar, Worker, Order, Invoice, Duty, Price, CustomsDuty, Excise, TransportCompany, \
    TransportCompanyPrice, Customer, PhotoGallery
from .forms import ParserForm, RegistrationForm, LoginForm, LogoutForm, OrderForm, OrderInOrdersForm, InvoiceForm, \
    NewInvoiceForm, DutyForm, PriceForm, CustomsDutyForm, ExciseForm, TransportCompanyForm, TransportCompanyPriceForm, \
    CustomerForm
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.mixins import LoginRequiredMixin
from datetime import date
from django.urls import reverse
from docx import Document
from docx.shared import Inches, RGBColor, Pt


#
# from django.http import HttpResponse
# from PyPDF2 import PdfFileReader, PdfFileWriter, PdfFileMerger
# import io
# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter


class HomePageView(LoginRequiredMixin, TemplateView):
    template_name = "home.html"

    def get(self, request, *args, **kwargs):
        if request.method == 'GET':
            return render(request, self.template_name)


class LoginPageView(TemplateView):
    template_name = "registration/login.html"

    def get(self, request, *args, **kwargs):
        form = LoginForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, username=username, password=password)
            # user = authenticate(request, **form.cleaned_data)
            print(user)
            if user is not None:
                login(request, user)
                # messages.success(request, "Вход выполнен")
                return redirect('home')
            messages.warning(request, "Неправильное имя пользователя или пароль")
        else:
            for field in form:
                print("Field Error:", field.name, field.errors)
            messages.error(request, "Некорректная форма")
        return render(request, 'registration/login.html', {'form': form})


def logout_view(request):
    logout(request)
    return redirect('home')


class RegistrationPageView(TemplateView):
    template_name = "registration/registration.html"

    def get(self, request, *args, **kwargs):
        if request.method == 'GET':
            form = RegistrationForm()
            form.fields['date_joined'].widget.attrs.update({'value': datetime.date.today()})
            form.fields['is_active'].widget.attrs.update({'value': True})
            form.fields['is_staff'].widget.attrs.update({'value': True})
            form.fields['is_superuser'].widget.attrs.update({'value': True})
            return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = RegistrationForm(request.POST)
            if form.is_valid():
                if form.username_clean() is None:
                    messages.warning(request, "Данное имя пользователя уже используется")
                    return render(request, 'registration/registration.html', {'form': form})
                if form.passport_clean() is None:
                    messages.warning(request, "Пользователь с таким паспортом уже существует")
                    return render(request, 'registration/registration.html', {'form': form})
                if form.clean_password2() is None:
                    messages.warning(request, "Пароли не совпадают")
                    return render(request, 'registration/registration.html', {'form': form})
                if form.save() == None:
                    messages.warning(request,
                                     "Этот номер паспорта уже используется. Пожалуйста, выберите другой номер паспорта.")
                    return render(request, 'registration/registration.html', {'form': form})
                else:
                    form.save()
                    messages.success(request, "Пользователь зарегистрирован")
            else:
                messages.error(request, "Некорректная форма")
        else:
            form = RegistrationForm()
        return render(request, 'registration/registration.html', {'form': form})


class GenreYear:
    """Жанры и года выхода фильмов"""

    def get_genres(self):
        print(Order.objects.all().values('order_status'))
        return Order.objects.all().values('order_status')

    def get_years(self):
        return Order.objects.all().values('order_status')


# class FilterView(GenreYear, ListView):
#     paginate_by = 5
#
#     def get_queryset(self):
#         queryset = Order.objects.filter(
#             Q(order_status__in=self.request.GET.getlist('genre'))
#         ).distinct()
#         print(queryset)
#         return queryset
#
#     def get_context_data(self, *args, **kwargs):
#         context = super().get_context_data(*args, **kwargs)
#         context["year"] = ''.join([f"year={x}&" for x in self.request.GET.getlist("year")])
#         context["genre"] = ''.join([f"genre={x}&" for x in self.request.GET.getlist("genre")])
#         print(context)
#         return context



class TransportCompanyPricePageView(TemplateView):
    template_name = 'transport_company_price.html'

    def get(self, request, *args, **kwargs):
        transport_company_price = TransportCompanyPrice.objects.get(pk=kwargs['price_id'])
        form = TransportCompanyPriceForm()
        form.fields['price'].initial = transport_company_price.price
        form.fields['place'].initial = transport_company_price.place
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = TransportCompanyPriceForm(request.POST)
            if form.is_valid():
                form.update(kwargs['price_id'])
                messages.info(request, "Цена ТК изменена")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, 'Ошибка валидации формы')
        tk_price = TransportCompanyPrice.objects.get(pk=kwargs['price_id'])
        return django.http.HttpResponseRedirect(reverse('transport_company_prices', kwargs={'transport_company_prices_id': tk_price.id_transport_company.pk}))




class TransportCompanyPricesNewPageView(TemplateView):
    template_name = 'transport_company_price.html'

    def get(self, request, *args, **kwargs):
        form = TransportCompanyPriceForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = TransportCompanyPriceForm(request.POST)
            if form.is_valid():
                print(kwargs['tk_id'])
                form.save(kwargs['tk_id'])
                messages.info(request, "Добавлена новая цена ТК")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, 'Ошибка валидации формы')
        return django.http.HttpResponseRedirect(reverse('transport_company_prices', kwargs={'transport_company_prices_id': kwargs['tk_id']}))


class TransportCompanyPricesPageView(TemplateView):
    template_name = 'transport_company_prices.html'

    def get(self, request, *args, **kwargs):
        transport_company_prices = TransportCompanyPrice.objects.filter(id_transport_company=kwargs['transport_company_prices_id'])
        return render(request, self.template_name, {'transport_company_prices': transport_company_prices})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST' and 'new' in request.POST:
            return django.http.HttpResponseRedirect(reverse('transport_company_prices_new', kwargs={'tk_id': kwargs['transport_company_prices_id']}))


class TransportCompaniesPageView(TemplateView):
    template_name = 'transport_companies.html'

    def get(self, request, *args, **kwargs):
        transport_companies = TransportCompany.objects.all()
        return render(request, self.template_name, {'transport_companies': transport_companies})

    def post(self, request, *args, **kwargs):
        if request.headers.get('x-requested-with') == 'XMLHttpRequest' and 'application' in request.POST:
            print(request.POST['number_contract'])
            document = docx.Document()
            styles = document.styles
            styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
            styles['Heading 2'].font.color.rgb = RGBColor(0, 0, 0)

            today_date = datetime.datetime.now().strftime('%d.%m.%Y')

            heading_text = f'ЗАЯВКА ОТ {today_date}г.'
            heading = document.add_heading(heading_text, 1)
            heading.alignment = 1
            heading.paragraph_format.space_after = Pt(0)
            heading.paragraph_format.space_before = Pt(0)

            heading_text = f'НА ОСНОВАНИИ ДОГОВОРА НА ОСУЩЕСТВЛЕНИЕ ПЕРЕВОЗКИ №' + request.POST['number_contract']
            heading = document.add_heading(heading_text, 1)
            heading.alignment = 1
            heading.paragraph_format.space_after = Pt(0)
            heading.paragraph_format.space_before = Pt(0)
            heading_text = f'ООО «Автолэнд ДВ» просит организовать доставку из г. Владивосток следующего груза:'
            heading = document.add_heading(heading_text, 2)
            heading.alignment = 1
            heading.paragraph_format.space_after = Pt(0)
            heading.paragraph_format.space_before = Pt(0)

            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Адрес получения'
            hdr_cells[1].text = 'Характер груза'
            hdr_cells[2].text = 'Грузополучатель'
            hdr_cells[3].text = '№ Дефектной ведомости'

            # Установка размера текста для заголовков таблицы
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Установите желаемый размер шрифта

            orders = Order.objects.all()
            records = []
            for order in orders:
                records.append({
                    'address': order.city,
                    'cargo': 'Легковой автомобиль: ' + str(order.id_car.title) + '\n' + 'Кузов: ' + str(order.id_car.the_body),
                    'customer': str(order.id_customer),
                    'dev_ved': str(order.pk),
                })

            for record in records:
                row_cells = table.add_row().cells
                row_cells[0].text = record['address']
                row_cells[1].text = record['cargo']
                row_cells[2].text = record['customer']
                row_cells[3].text = record['dev_ved']
                # Установка размера текста для ячеек таблицы
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)  # Установите желаемый размер шрифта

            file_path = 'media/client_contract/demo.docx'
            document.save(file_path)

            if os.path.exists(file_path):
                file_url = f'/media/client_contract/demo.docx'
                return JsonResponse({'file_url': file_url})

        return JsonResponse({'error': 'Invalid request'}, status=400)
class TransportCompanyNewPageView(TemplateView):
    template_name = 'transport_company.html'

    def get(self, request, *args, **kwargs):
        print('NEW')
        form = TransportCompanyForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = TransportCompanyForm(request.POST, request.FILES)

            if form.is_valid():
                form.save()
                messages.info(request, "Добавлена новая транспортная компания")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
        transport_companies = TransportCompany.objects.all()
        return render(request, 'transport_companies.html', {'transport_companies': transport_companies})


class TransportCompanyPageView(TemplateView):
    template_name = 'transport_company.html'

    def get(self, request, *args, **kwargs):
        form = TransportCompanyForm()
        # form.fields['name'].initial
        transport_company = TransportCompany.objects.get(id=kwargs['transport_company_id'])
        form.fields['title'].initial = transport_company.title
        form.fields['number_contract'].initial = transport_company.number_contract
        form.fields['contract'].initial = transport_company.contract
        form.fields['contract'].widget.input_text = 'Заменить'
        form.fields['contract'].widget.clear_checkbox_label = ''
        form.fields['contract'].widget.initial_text = ''
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST' and 'save' in request.POST:
            form = TransportCompanyForm(request.POST, request.FILES)
            if form.is_valid():

                transport_company = TransportCompany.objects.get(id=kwargs['transport_company_id'])
                if transport_company.contract == '':
                    contract = request.FILES['contract']
                else:
                    contract = transport_company.contract
                form.update(kwargs['transport_company_id'], contract)
                messages.info(request, "Транспортная компания изменена")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
            transport_companies = TransportCompany.objects.all()
            return render(request, 'transport_companies.html', {'transport_companies': transport_companies})

        if request.method == 'POST' and 'new_price' in request.POST:
            transport_company_prices = TransportCompanyPrice.objects.filter(
                id_transport_company=kwargs['transport_company_id'])
            transport_company_prices = {
                'transport_company_prices': transport_company_prices
            }
            print(kwargs['transport_company_id'])
            return django.http.HttpResponseRedirect(reverse('transport_company_prices', kwargs={'transport_company_prices_id': kwargs['transport_company_id']}))
            # return render(request, 'transport_company_prices.html', transport_company_prices)


class CustomsDutysPageView(TemplateView):
    template_name = 'customs_dutys.html'

    def get(self, request, *args, **kwargs):
        customs_dutys = CustomsDuty.objects.all()
        return render(request, self.template_name, {'customs_dutys': customs_dutys})


class CustomsDutyNewPageView(TemplateView):
    template_name = 'customs_duty.html'

    def get(self, request, *args, **kwargs):
        print('NEW')
        form = CustomsDutyForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = CustomsDutyForm(request.POST)
            if form.is_valid():
                form.save()
                messages.info(request, "Добавлена новая таможенная пошлина")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
        customs_dutys = CustomsDuty.objects.all()
        return render(request, 'customs_dutys.html', {'customs_dutys': customs_dutys})


class CustomsDutyPageView(TemplateView):
    template_name = 'customs_duty.html'

    def get(self, request, *args, **kwargs):
        form = CustomsDutyForm()
        # form.fields['name'].initial
        print('Просим')
        customs_duty = CustomsDuty.objects.get(id=kwargs['customs_duty_id'])
        form.fields['type'].initial = customs_duty.type
        form.fields['value_first'].initial = customs_duty.value_first
        form.fields['value_last'].initial = customs_duty.value_last
        form.fields['bet'].initial = customs_duty.bet
        form.fields['date_of_action'].initial = customs_duty.date_of_action
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = CustomsDutyForm(request.POST)
            if form.is_valid():
                print(kwargs['customs_duty_id'])

                form.update(kwargs['customs_duty_id'])
                messages.info(request, "Таможенная пошлина изменена")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
        customs_dutys = CustomsDuty.objects.all()
        return render(request, 'customs_dutys.html', {'customs_dutys': customs_dutys})


class ExcisesPageView(TemplateView):
    template_name = 'excises.html'

    def get(self, request, *args, **kwargs):
        excises = Excise.objects.all()
        return render(request, self.template_name, {'excises': excises})


class ExciseNewPageView(TemplateView):
    template_name = 'excise.html'

    def get(self, request, *args, **kwargs):
        print('NEW')
        form = ExciseForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = ExciseForm(request.POST)
            if form.is_valid():
                form.save()
                messages.info(request, "Добавлен новый акциз")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
        excises = Excise.objects.all()
        return render(request, 'excises.html', {'excises': excises})


class ExcisePageView(TemplateView):
    template_name = 'excise.html'

    def get(self, request, *args, **kwargs):
        form = ExciseForm()
        # form.fields['name'].initial
        excise = Excise.objects.get(id=kwargs['excise_id'])
        form.fields['power_first_car'].initial = excise.power_first_car
        form.fields['power_last_car'].initial = excise.power_last_car
        form.fields['bet'].initial = excise.bet
        form.fields['date_of_action'].initial = excise.date_of_action
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = ExciseForm(request.POST)
            if form.is_valid():
                print(kwargs['excise_id'])

                form.update(kwargs['excise_id'])
                messages.info(request, "Акциз изменен")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
        excises = Excise.objects.all()
        return render(request, 'excises.html', {'excises': excises})


class PricesPageView(TemplateView):
    template_name = 'prices.html'

    def get(self, request, *args, **kwargs):
        prices = Price.objects.all()
        return render(request, self.template_name, {'prices': prices})


class PriceNewPageView(TemplateView):
    template_name = 'price.html'

    def get(self, request, *args, **kwargs):
        print('NEW')
        form = PriceForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = PriceForm(request.POST)
            if form.is_valid():
                form.save()
                messages.info(request, "Добавлена новая цена перевозки авто из ЯП")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
        prices = Price.objects.all()
        return render(request, 'prices.html', {'prices': prices})


class PricePageView(TemplateView):
    template_name = 'price.html'

    def get(self, request, *args, **kwargs):
        form = PriceForm()
        # form.fields['name'].initial
        price = Price.objects.get(id=kwargs['price_id'])
        form.fields['price_first_car'].initial = price.price_first_car
        if price.price_last_car is not None:
            form.fields['price_last_car'].initial = price.price_last_car
        # else:
        #     form.fields['']
        form.fields['price_transportation'].initial = price.price_transportation
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = PriceForm(request.POST)
            if form.is_valid():
                print(kwargs['price_id'])

                form.update(kwargs['price_id'])
                messages.info(request, "Цена перевоза авто из ЯП изменена")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
        prices = Price.objects.all()
        return render(request, 'prices.html', {'prices': prices})


class DutiesPageView(TemplateView):
    template_name = 'duties.html'

    def get(self, request, *args, **kwargs):
        duties = Duty.objects.all()
        return render(request, self.template_name, {'duties': duties})


class DutyNewPageView(TemplateView):
    template_name = 'duty.html'

    def get(self, request, *args, **kwargs):
        form = DutyForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = DutyForm(request.POST)
            if form.is_valid():
                form.save()
                messages.info(request, "Добавлена новая ставка утилизационного сбора")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
        duties = Duty.objects.all()
        return render(request, 'duties.html', {'duties': duties})


class DutyPageView(TemplateView):
    template_name = 'duty.html'

    def get(self, request, *args, **kwargs):
        form = DutyForm()
        # form.fields['name'].initial
        duty = Duty.objects.get(id=kwargs['duty_id'])
        form.fields['volume_first'].initial = duty.volume_first
        form.fields['volume_last'].initial = duty.volume_last
        form.fields['coefficient_less_3'].initial = duty.coefficient_less_3
        form.fields['coefficient_more_3'].initial = duty.coefficient_more_3
        form.fields['date_of_action'].initial = duty.date_of_action
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = DutyForm(request.POST)
            if form.is_valid():
                print(kwargs['duty_id'])

                form.update(kwargs['duty_id'])
                messages.info(request, "Ставка утилизационного сбора изменена")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.info(request, "Ошибка валидации формы")
        duties = Duty.objects.all()
        return render(request, 'duties.html', {'duties': duties})


class WorkersPageView(TemplateView):
    template_name = "workers.html"

    def get(self, request, *args, **kwargs):
        workers = Worker.objects.filter(is_superuser=False)
        return render(request, 'workers.html', {'workers': workers})


class WorkersCardPageView(TemplateView):
    template_name = "worker_card.html"

    def get(self, request, *args, **kwargs):
        worker = Worker.objects.get(id=kwargs.get('worker_id'))
        worker_data = Worker.objects.all()
        form = RegistrationForm()
        form.fields['username'].widget.attrs.update({'value': worker.username})
        form.fields['full_name'].widget.attrs.update({'value': worker.full_name})
        # Вот тут хуй знает как сделать не нашел
        form.fields['job_title'].widget.attrs.update({'value': worker.job_title})
        # Вот тут хуй знает как сделать не нашел
        form.fields['passport'].widget.attrs.update({'value': worker.passport})
        form.fields['phone_number'].widget.attrs.update({'value': worker.phone_number})
        form.fields['password'].widget.attrs.update({'value': worker.password})
        form.fields['password2'].widget.attrs.update({'value': worker.password})

        return render(request, 'worker_card.html', {'worker': worker, 'worker_data': worker_data, 'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = RegistrationForm(request.POST)
            if form.is_valid():
                # form.username_clean()
                if form.clean_password2() is None:
                    # message = messages.info(request, 'Your password has been changed successfully!')
                    messages.warning(request, "Пароли не совпадают")
                    return render(request, 'registration/registration.html', {'form': form})
                if form.passport_clean() is None:
                    # message = messages.info(request, 'Your password has been changed successfully!')
                    messages.warning(request, "Пользователь с таким паспортом уже существует")
                    return render(request, 'registration/registration.html', {'form': form})
                if form.username_clean() is None:
                    # message = messages.info(request, 'Your password has been changed successfully!')
                    messages.warning(request, "Пользователь с таким именем уже существует")
                    return render(request, 'registration/registration.html', {'form': form})
                form.update()
                messages.success(request, "Данные обновлены")
            else:
                for field in form:
                    print("Field Error:", field.name, field.errors)
                messages.error(request, "Некорректная форма")

        else:
            form = RegistrationForm()
        return render(request, 'registration/registration.html', {'form': RegistrationForm()})


class OrderInOrdersPageView(TemplateView):
    template_name = 'order_in_orders.html'

    def create_dev_ved(self, order_id):
        order = Order.objects.get(pk=order_id)
        document = docx.Document()
        styles = document.styles
        styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        heading = 'Дефектная ведомость №' + ' ' + str(order.pk)
        heading = document.add_heading(heading, 1)
        heading.alignment = 1
        para = 'Дата: ' + str(date.today())
        para = document.add_paragraph(para)
        para.paragraph_format.space_after = Inches(0.001)
        para = document.add_paragraph('Агент:')
        para.paragraph_format.space_after = Inches(0.001)
        para = document.add_paragraph('ВладивостокМоторс')
        # para.alignment = 2
        para.paragraph_format.space_after = Inches(0.001)
        para = document.add_paragraph('г. Владивосток,')
        # para.alignment = 2
        para.paragraph_format.space_after = Inches(0.001)
        # para.alignment = 2
        para = document.add_paragraph('ул. Авроровская 19А, к. 195')
        para.paragraph_format.space_after = Inches(0.001)
        # para.alignment = 2
        para = document.add_paragraph('тел: +7908237482')
        para.paragraph_format.space_after = Inches(0.001)
        # para.alignment = 2
        para = document.add_paragraph('e-mail: vladmotors@vladmotors.ru')
        para.paragraph_format.space_after = Inches(0.001)
        # para.alignment = 2
        document.add_picture('media/def_ved_img.png', width=Inches(7))

        records = (
            ('Модель', order.id_car.title),
            ('№ Кузова', order.id_car.the_body),
            ('Цвет', order.id_car.color)
        )

        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Наименование'
        for qty in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty[0])
            row_cells[1].text = str(qty[1])
        title = os.path.join(settings.MEDIA_ROOT, 'client_contract/', f'Дефектная_ведомость_{order.id_customer.first_name_client[0]}_{order.id_customer.patronymic_client[0]}_{order.id_customer.last_name_client}.docx')
        document.save(title)
        context = {

            'title': title
        }
        return context

    def create_contract(self, order_id):
        order = Order.objects.get(pk=order_id)
        document = docx.Document()
        styles = document.styles
        styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)

        heading = 'Агентский договор № ' + str(order.pk) + '/' + str(order.id_customer.pk)
        heading = document.add_heading(heading, 1)
        heading.alignment = 1

        prim1 = document.add_paragraph('(на приобретение транспортного средства, его доставку в РФ и оформление)')
        prim1.alignment = 1

        document.add_paragraph(f'г. Владивосток \t\t\t\t\t\t\t                    {str(date.today())}')
        paragraph1 = document.add_paragraph(
            f'Общество с ограниченной ответственностью АвтолендДВ, именуемое в тексте договора "Поставщик", в лице {order.id_worker.full_name}, действующего на основании Устава с одной стороны, и {order.id_customer.last_name_client} {order.id_customer.first_name_client} {order.id_customer.patronymic_client}, дата рождения {order.id_customer.date_of_birth}г, паспорт {order.id_customer.passport_series}№{order.id_customer.passport_number}, выдан {order.id_customer.passport_department_name}, код подразделения {order.id_customer.passport_department_code}, дата выдачи {order.id_customer.date_of_issue}г, зарегистрирован: {order.id_customer.address},именуемый в тексте договора "Заказчик", с другой стороны, заключили настоящий договор о нижеследующем:')
        paragraph1.paragraph_format.first_line_indent = Inches(0.5)
        paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_heading('1. Предмет договора', 1)

        document.add_paragraph(
            '1.1. Поставщик обязуется за вознаграждение совершать по поручению Заказчика юридические и иные действия от своего имени, но за счет Заказчика, либо от имени и за счет Заказчика.')
        document.add_paragraph(
            '1.2. Поставщик приобретает права и становится обязанным по сделке, совершенной с третьим лицом от своего имени за счет Заказчика.')
        document.add_paragraph(
            '1.3. По сделке, совершенной Поставщиком с третьим лицом от имени и за счет Заказчика, права и обязанности возникают у Заказчика.')
        document.add_paragraph(
            '1.4. В соответствии с настоящим договором Поставщик обязуется по поручению Заказчика организовать покупку транспортного средства (далее по тексту ТС) на автомобильных аукционах в Японии и доставку указанного ТС до места получения ТС в соответствии с заявкой (поручением) Заказчика.')
        document.add_paragraph(
            '1.5. Для исполнения поручения Заказчика Поставщик обязуется совершить следующие действия:')
        document.add_paragraph('- осуществить покупку указанного Заказчиком ТС на аукционе Японии;')
        document.add_paragraph('- осуществить доставку приобретенного ТС в порт погрузки в Японии;')
        document.add_paragraph('- осуществить доставку приобретенного ТС морским транспортом до порта г. Владивосток;')
        document.add_paragraph(
            '- осуществить действия по таможенной очистке ТС в г. Владивосток, в том числе оформить необходимые таможенные документы;')
        document.add_paragraph('- осуществить передачу приобретенного ТС Заказчику.')
        document.add_paragraph(
            '1.6. Для осуществления действий указанных в п.1.5. настоящего договора Поставщик заключает от своего имени необходимые договоры, в том числе агентские, подписывает необходимые документы, а также производит необходимые платежи.')

        title = os.path.join(settings.MEDIA_ROOT, 'client_contract/', f'Договор_с_клиентом_{order.id_customer.first_name_client[0]}_{order.id_customer.patronymic_client[0]}_{order.id_customer.last_name_client}.docx')
        document.save(title)
        context = {

            'title': title
        }
        return context

    def get(self, request, *args, **kwargs):
        order = Order.objects.get(id_order=kwargs['order_id'])
        car = order.id_car
        photos = order.photos.all()
        photo = PhotoCar.objects.filter(id_car=car.pk)[:1][0].photo
        form = OrderInOrdersForm()
        file_fields = [
            'sbts', 'ptd', 'client_contract', 'def_ved', 'export_certificate',
            'consignment', 'received_ptd', 'invoice', 'payment_order', 'contract_japan', 'photos'
        ]
        form.fields['export_certificate_number'].initial = order.export_certificate_number
        form.fields['id_order'].widget.attrs.update({'value': order.id_order})
        form.fields['first_name_client'].widget.attrs.update({'value': order.id_customer.first_name_client})
        form.fields['last_name_client'].widget.attrs.update({'value': order.id_customer.last_name_client})
        form.fields['patronymic_client'].widget.attrs.update({'value': order.id_customer.patronymic_client})
        form.fields['telephone'].widget.attrs.update({'value': order.id_customer.telephone})
        form.fields['date_start'].widget.attrs.update({'value': order.date_start})
        if order.price is not None:
            form.fields['price'].widget.attrs.update({'value': str(order.price)})
        for field_name in file_fields:
            form.fields[field_name].widget.initial_text = ''
            form.fields[field_name].widget.input_text = 'Заменить'
            form.fields[field_name].widget.clear_checkbox_label = ''
        if order.date_end is not None:
            form.fields['date_end'].widget.attrs.update({'value': order.date_end, 'readonly': 'True'})
        if order.comment is not None:
            form.fields['comment'].initial = order.comment
        if order.consignment is not None:
            form.fields['consignment'].initial = order.consignment
        if order.sbts is not None:
            form.fields['sbts'].initial = order.sbts
        if order.ptd is not None:
            form.fields['ptd'].initial = order.ptd
        if order.contract is not None:
            form.fields['client_contract'].initial = order.contract
        if order.defective_statement is not None:
            form.fields['def_ved'].initial = order.defective_statement
        if order.export_certificate is not None:
            form.fields['export_certificate'].initial = order.export_certificate
        if order.received_ptd is not None:
            form.fields['received_ptd'].initial = order.received_ptd
        if order.invoice is not None:
            form.fields['invoice'].initial = order.invoice
        if order.payment_order is not None:
            form.fields['payment_order'].initial = order.payment_order
        user_id = request.user.id
        user = Worker.objects.get(pk=user_id)
        if user.job_title == 'Клиент':
            for field in form.fields.values():
                field.widget.attrs['readonly'] = True
            for field_name in file_fields:
                form.fields[field_name].widget.attrs['disabled'] = True
        return render(request, self.template_name, {'order': order, 'form': form, 'order_id': order.id_order, 'car': car, 'photo': photo,  'photos': photos})

    def post(self, request, *args, **kwargs):
        print(request.FILES)
        if request.method == 'POST' and 'update' in request.POST:
            form = OrderInOrdersForm(request.POST, request.FILES)
            if form.is_valid():

                order = Order.objects.filter(id_order=form.cleaned_data['id_order'])
                order = order[0]
                # is_initial
                print(form.initial)
                print(form.cleaned_data['date_end'])
                # print(form.fields['ptd'].initial)
                print(request.FILES.get('ptd'))

                if 'ptd' in request.FILES:
                    order.ptd = request.FILES.get('ptd')
                else:
                    order.ptd = order.ptd

                if 'sbts' in request.FILES:
                    order.sbts = request.FILES.get('sbts')
                else:
                    order.sbts = order.sbts

                if 'client_contract' in request.FILES:
                    order.contract = request.FILES.get('client_contract')
                else:
                    order.contract = order.contract

                if 'def_ved' in request.FILES:
                    order.defective_statement = request.FILES.get('def_ved')
                else:
                    order.defective_statement = order.defective_statement

                if 'export_certificate' in request.FILES:
                    order.export_certificate = request.FILES.get('export_certificate')
                else:
                    order.export_certificate = order.export_certificate

                if 'consignment' in request.FILES:
                    order.consignment = request.FILES.get('consignment')
                else:
                    order.consignment = order.consignment

                if 'received_ptd' in request.FILES:
                    order.received_ptd = request.FILES.get('received_ptd')
                else:
                    order.received_ptd = order.received_ptd

                if 'invoice' in request.FILES:
                    order.invoice = request.FILES.get('invoice')
                else:
                    order.invoice = order.invoice

                if 'payment_order' in request.FILES:
                    order.payment_order = request.FILES.get('payment_order')
                else:
                    order.payment_order = order.payment_order

                print(order.ptd)
                if order.ptd and order.sbts:
                    order.order_status = order.WAITING_TO_BE_SENT


                messages.success(request, "Заказ изменен")
                form.save()
            else:
                messages.error(request, "Некорректная форма")
                for field in form:
                    print("Field Error:", field.name, field.errors)
            return django.http.HttpResponseRedirect(reverse('orders'))

        elif request.method == 'POST' and 'calculate_price' in request.POST:
            form = OrderInOrdersForm(request.POST, request.FILES)
            print('считаем цену')
            order = Order

            if form.is_valid():
                order = Order.objects.get(id_order=form.cleaned_data['id_order'])
                price = form.cleaned_data['price_for_buhgalter']
                power = int(form.cleaned_data['power'])
                duties = Duty.objects.all()
                prices = Price.objects.all()
                excises = Excise.objects.all()
                customs_dutys = CustomsDuty.objects.all()
                volume = int(order.id_car.volume) / 1000
                year = int(datetime.datetime.now().year) - int(order.id_car.year_car)
                base_bet = 20000
                coefficient_bet = 0
                coefficient_excise = 0
                coefficient_customs_duty = 0
                price_transportation = 0
                for duty in duties:
                    if duty.volume_first <= volume <= duty.volume_last and year > 3:
                        coefficient_bet = duty.coefficient_more_3
                    elif duty.volume_first <= volume <= duty.volume_last and year <= 3:
                        coefficient_bet = duty.coefficient_less_3
                #
                for price_data in prices:
                    if price_data.price_first_car <= int(price) < price_data.price_last_car:
                        price_transportation = price_data.price_transportation
                    elif price_data.price_last_car == 0 and price_transportation == 0:
                        price_transportation = price_data.price_transportation
                volume_or_price = 0
                if year < 3:
                    customs_dutys = CustomsDuty.objects.filter(type=CustomsDuty.TYPE_CHOICE[0][1])
                    volume_or_price = price
                elif 3 <= year < 5:
                    customs_dutys = CustomsDuty.objects.filter(type=CustomsDuty.TYPE_CHOICE[1][1])
                    volume_or_price = volume * 1000
                else:
                    customs_dutys = CustomsDuty.objects.filter(type=CustomsDuty.TYPE_CHOICE[2][1])
                    volume_or_price = volume * 1000

                for customs_duty in customs_dutys:
                    if customs_duty.value_first <= int(volume_or_price) <= customs_duty.value_last:
                        coefficient_customs_duty = int(volume) * 1000 * customs_duty.bet
                        print(volume * 1000, year, customs_duty.bet)
                    elif customs_duty.value_last == 0 and coefficient_customs_duty == 0:
                        coefficient_customs_duty = int(volume) * 1000 * customs_duty.bet

                for excise in excises:
                    if excise.power_first_car <= power <= excise.power_last_car:
                        coefficient_excise = power * excise.bet
                        print(coefficient_excise)
                    elif coefficient_excise == 0 and excise.power_last_car == 0:
                        coefficient_excise = power * excise.bet

                # Ставки утилизационного сбора
                final_price = int(price) + base_bet * coefficient_bet
                # Цена доставки тачки из Японии
                final_price = final_price + price_transportation
                # Таможенная пошлина customs_dutys
                final_price = final_price + coefficient_customs_duty
                # Акциз excises
                final_price = final_price + coefficient_excise
                # НДС (стоимость авто+таможенная пошлина+акциз)*20%
                nds = (int(price) + int(coefficient_customs_duty) + int(coefficient_excise)) * 0.2
                final_price = final_price + int(nds)
                print(power)
                print(int(price), base_bet * coefficient_bet, price_transportation, coefficient_excise,
                      coefficient_customs_duty, nds)
                print(final_price)
                form.fields['price'].widget.attrs.update({'value': final_price})


            return render(request, 'order_in_orders.html', {'form': form, 'order': order})

        elif request.method == 'POST' and 'calculate_price' in request.POST:
            return render(request, 'order_in_orders.html')

        elif request.method == 'POST' and 'create_contract' in request.POST:
            form = OrderInOrdersForm(request.POST, request.FILES)
            if form.is_valid():
                order = Order.objects.get(id_order=form.cleaned_data['id_order'])

                title = self.create_contract(order.pk)['title']
                file_path = title
                if os.path.exists(file_path):
                    with open(file_path, 'rb') as fh:
                        response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")

                        # Кодирование имени файла для использования в заголовке Content-Disposition
                        filename = os.path.basename(file_path)
                        filename_header = quote(filename)

                        response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{filename_header}'
                        return response

        elif request.method == 'POST' and 'create_defective_statement' in request.POST:
            form = OrderInOrdersForm(request.POST, request.FILES)

            if form.is_valid():
                order = Order.objects.get(id_order=form.cleaned_data['id_order'])

                title = self.create_dev_ved(order.pk)['title']
                print(title)
                file_path = title
                if os.path.exists(file_path):
                    with open(file_path, 'rb') as fh:
                        response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")

                        # Кодирование имени файла для использования в заголовке Content-Disposition
                        filename = os.path.basename(file_path)
                        filename_header = quote(filename)

                        response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{filename_header}'
                        return response

        user_id = request.user.id
        orders = Order.objects.filter(date_end=None, id_worker=user_id)
        return render(request, 'orders.html', {"orders": orders})


class OrdersPageView(TemplateView):
    template_name = "orders.html"

    def get(self, request, *args, **kwargs):
        user_id = request.user.id
        user = Worker.objects.get(pk=user_id)
        if user.job_title == 'Менеджер':
            orders = Order.objects.filter(date_end=None, id_worker=user_id)
        elif user.job_title == 'Оперативник':
            orders = Order.objects.all()
        elif user.job_title == 'Клиент':
            orders = Order.objects.filter(date_end=None, id_worker=user_id)
        else:
            orders = Order.objects.all()
        return render(request, 'orders.html', {'orders': orders})

    def post(self, request, *args, **kwargs):
        if request.headers.get('x-requested-with') == 'XMLHttpRequest' and 'auc_doc_btn' in request.POST:
            document = docx.Document()
            styles = document.styles
            styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)

            heading = document.add_heading('Список автомобилей к покупке на аукционе', 1)
            heading.alignment = 1



            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Наименование аукциона'
            hdr_cells[1].text = 'Номер лота'
            hdr_cells[2].text = 'Диапазон бюджета'

            # Установка размера текста для заголовков таблицы
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Установите желаемый размер шрифта
            orders = Order.objects.all()
            records = []
            for order in orders:
                records.append({
                    'auction_name': order.id_car.auc_name,
                    'lot_number': order.id_car.auc_number,
                    'budget_range': str(order.id_car.price) + ' - ' + str(order.price_customer)
                })

            for record in records:
                row_cells = table.add_row().cells
                row_cells[0].text = record['auction_name']
                row_cells[1].text = record['lot_number']
                row_cells[2].text = record['budget_range']
                # Установка размера текста для ячеек таблицы
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)  # Установите желаемый размер шрифта

            file_path = 'media/client_contract/demo.docx'
            document.save(file_path)

            if os.path.exists(file_path):
                file_url = f'/media/client_contract/demo.docx'
                return JsonResponse({'file_url': file_url})

        if request.headers.get('x-requested-with') == 'XMLHttpRequest' and 'trans_btn' in request.POST:
            document = docx.Document()
            styles = document.styles
            styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
            styles['Heading 2'].font.color.rgb = RGBColor(0, 0, 0)

            today_date = datetime.datetime.now().strftime('%d.%m.%Y')

            heading_text = f'ЗАЯВКА ОТ {today_date}г.'
            heading = document.add_heading(heading_text, 1)
            heading.alignment = 1
            heading.paragraph_format.space_after = Pt(0)
            heading.paragraph_format.space_before = Pt(0)

            heading_text = f'НА ОСНОВАНИИ ДОГОВОРА НА ОСУЩЕСТВЛЕНИЕ ПЕРЕВОЗКИ №1'
            heading = document.add_heading(heading_text, 1)
            heading.alignment = 1
            heading.paragraph_format.space_after = Pt(0)
            heading.paragraph_format.space_before = Pt(0)
            heading_text = f'ООО «Автолэнд ДВ» просит организовать доставку в г. Владивосток следующего груза:'
            heading = document.add_heading(heading_text, 2)
            heading.alignment = 1
            heading.paragraph_format.space_after = Pt(0)
            heading.paragraph_format.space_before = Pt(0)

            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Характер груза'
            hdr_cells[1].text = 'Экспортный серктификат'


            # Установка размера текста для заголовков таблицы
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Установите желаемый размер шрифта
            orders = Order.objects.all()
            records = []
            for order in orders:
                records.append({
                    'auction_name': 'Легковой автомобиль: ' + str(order.id_car.title) + '\n' + 'Кузов: ' + str(order.id_car.the_body),
                    'lot_number': str(order.export_certificate_number),
                })

            for record in records:
                row_cells = table.add_row().cells
                row_cells[0].text = record['auction_name']
                row_cells[1].text = record['lot_number']
                # Установка размера текста для ячеек таблицы
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)  # Установите желаемый размер шрифта




            file_path = 'media/client_contract/demo.docx'
            document.save(file_path)

            if os.path.exists(file_path):
                file_url = f'/media/client_contract/demo.docx'
                return JsonResponse({'file_url': file_url})

        return JsonResponse({'error': 'Неверный запрос'}, status=400)


class OrderPageView(TemplateView):
    template_name = "order.html"

    def get(self, request, *args, **kwargs):
        print(TransportCompanyPrice.objects.values_list('place', flat=True).distinct())
        car = Car.objects.get(id_car=kwargs.get('car_id'))
        form = OrderForm()

        photo = PhotoCar.objects.filter(id_car=kwargs.get('car_id'))[:1][0].photo
        form.fields['id_car'].widget.attrs.update({'value': car.id_car})
        user_name = Worker.objects.filter(id=request.user.id)[0]
        form.fields['worker'].widget.attrs.update({'value': user_name})
        form.fields['price'].initial = str(car.price) + ' р.'
        print(kwargs.get('customer_id'))
        form.fields['customer'].initial = kwargs.get('customer_id')
        return render(request, 'order.html', {'car': car, 'form': form, 'photo': photo})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST' and 'create' in request.POST:
            form = OrderForm(request.POST)
            if form.is_valid():
                form.save()
                messages.success(request, "Заказ создан")
                return django.http.HttpResponseRedirect(reverse('orders'))
            else:
                messages.error(request, "Некорректная форма")
                return render(request, 'order.html', {'form': form})
        elif request.method == 'POST' and 'customer' in request.POST:
            return django.http.HttpResponseRedirect(reverse('customer_new_for_order', kwargs={'car_id': kwargs.get('car_id')}))

        user_id = request.user.id
        orders = Order.objects.filter(date_end=None, id_worker=user_id)
        return render(request, 'orders.html', {'orders': orders})


class CustomerPageView(TemplateView):
    template_name = 'customer.html'

    def get(self, request, *args, **kwargs):
        form = CustomerForm()
        customer = Customer.objects.get(pk=kwargs['customer_id'])
        form.fields['first_name_client'].initial = customer.first_name_client
        form.fields['last_name_client'].initial = customer.last_name_client
        form.fields['patronymic_client'].initial = customer.patronymic_client
        form.fields['date_of_birth'].initial = customer.date_of_birth
        form.fields['place_of_birth'].initial = customer.place_of_birth
        form.fields['passport_number'].initial = customer.passport_number
        form.fields['passport_series'].initial = customer.passport_series
        form.fields['passport_department_code'].initial = customer.passport_department_code
        form.fields['passport_department_name'].initial = customer.passport_department_name
        form.fields['telephone'].initial = customer.telephone
        form.fields['address'].initial = customer.address
        form.fields['date_of_issue'].initial = customer.date_of_issue
        form.fields['inn'].widget.initial_text = ''
        form.fields['inn'].widget.input_text = 'Заменить'
        form.fields['inn'].widget.clear_checkbox_label = ''
        if customer.inn is not None:
            form.fields['inn'].initial = customer.inn
        form.fields['passport'].widget.initial_text = ''
        form.fields['passport'].widget.input_text = 'Заменить'
        form.fields['passport'].widget.clear_checkbox_label = ''
        if customer.passport is not None:
            form.fields['passport'].initial = customer.passport
        form.fields['registration'].widget.initial_text = ''
        form.fields['registration'].widget.input_text = 'Заменить'
        form.fields['registration'].widget.clear_checkbox_label = ''
        if customer.registration is not None:
            form.fields['registration'].initial = customer.registration


        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = CustomerForm(request.POST, request.FILES)
            if form.is_valid():
                customer = get_object_or_404(Customer, pk=kwargs['customer_id'])
                if 'inn' in request.FILES:
                    customer.inn = request.FILES.get('inn')
                else:
                    customer.inn = customer.inn
                if 'passport' in request.FILES:
                    customer.passport = request.FILES.get('passport')
                else:
                    customer.passport = customer.passport
                if 'registration' in request.FILES:
                    customer.registration = request.FILES.get('registration')
                else:
                    customer.registration = customer.registration
                customer.save()
                form.update_customer(kwargs['customer_id'])
                messages.success(request, "Клиент изменен")

                return django.http.HttpResponseRedirect(reverse('customers'))
            else:
                messages.error(request, "Некорректная форма")
                for field in form:
                    print("Field Error:", field.name, field.errors)
                return render(request, 'customer.html', {'form': form})
        else:
            form = CustomerForm()
        customers = Customer.objects.all()
        return render(request, 'customers.html', {'customers': customers})


class CustomersPageView(TemplateView):
    template_name = 'customers.html'

    def get(self, request, *args, **kwargs):
        customers = Customer.objects.all()
        return render(request, 'customers.html', {'customers': customers})



class CustomerNewPageView(TemplateView):
    template_name = 'customer.html'

    def get(self, request, *args, **kwargs):
        form = CustomerForm()
        if kwargs.get('car_id'):
            print('ЕСТЬ')
            car = Car.objects.get(id_car=kwargs.get('car_id'))
            return render(request, self.template_name, {'form': form, 'car': car})
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            form = CustomerForm(request.POST)
            if form.is_valid():
                if Customer.objects.filter(passport_number=form.cleaned_data['passport_number'], passport_series=form.cleaned_data['passport_series']).exists():
                    print('Такой уже есть')
                    messages.error(request, "Такой уже есть")
                    return django.http.HttpResponseRedirect(
                        reverse('customer_new_for_order', kwargs={'car_id': kwargs.get('car_id')}))

                customer = form.save()
                messages.success(request, "Клиент создан")
                if 'save_and_continue' in request.POST:
                    return django.http.HttpResponseRedirect(reverse('order_with_customer', kwargs={'car_id': kwargs.get('car_id'), 'customer_id': customer}))
                return django.http.HttpResponseRedirect(reverse('customers'))
            else:
                messages.error(request, "Некорректная форма")
                for field in form:
                    print("Field Error:", field.name, field.errors)
                return render(request, 'customer.html', {'form': form})
        else:
            form = CustomerForm()
        customers = Customer.objects.all()
        return render(request, 'customers.html', {'customers': customers})


class CatalogPageView(TemplateView):

    def get(self, request, *args, **kwargs):
        cars = Car.objects.all()
        cars = Car.objects.filter(auc_date__gte=datetime.date.today())

        for el in cars:
            el.image = PhotoCar.objects.filter(id_car=el.id_car)[:1][0].photo

        return render(request, 'catalog.html', {'cars': cars})


class CarPageView(TemplateView):
    template_name = "car.html"

    def get(self, request, *args, **kwargs):
        car = Car.objects.get(id_car=kwargs.get('car_id'))
        photo = PhotoCar.objects.filter(id_car=car)

        # Достаем данные из excel
        file_path = 'cars_price.xlsx'
        workbook = load_workbook(filename=file_path)

        models_list = workbook.worksheets[1]
        mark_list = workbook.worksheets[0]

        models = []
        skip_first_row = False
        for row in models_list.iter_rows(values_only=True, min_row=2 if skip_first_row else 1):
            model = {
                'id': row[0],
                'mark': row[1],
                'model': row[2],
                'price': row[3],
            }
            models.append(model)
        marks = []
        for row in mark_list.iter_rows(values_only=True, min_row=2 if skip_first_row else 1):
            mark = {
                'id': row[0],
                'mark': row[1]
            }
            marks.append(mark)
        # Достаем данные из excel

        # Узнаем цену машины
        car_for_test = Car.objects.all()
        for car_test in car_for_test:
            car_title = car_test.title.split()
            is_car = False
            price = 0
            for car_for_test in models:

                # if str(marks[int(car_for_test['mark'])-1]['mark']) == 'Acura':
                #     print(car_for_test, marks[int(car_for_test['mark'])-1]['mark'], 'true', int(car_for_test['mark']))
                # print(marks[int(car_for_test['mark'])-1]['mark'], car_title[0].lower())
                if (str(car_for_test['model']).lower() == car_title[1].lower()
                        and
                        str(marks[int(car_for_test['mark']) - 1]['mark']).lower() == car_title[0].lower()):
                    is_car = True
                    price = car_for_test['price']
            if is_car is False:
                price = 'Нет информации о цене'

            print(car_test.title, price)

        return render(request, 'car.html', {'car': car, 'photo': photo})


class ParserPageView(TemplateView):
    template_name = "parser.html"

    def get(self, request, *args, **kwargs):
        form = ParserForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        linked_list = list()
        if request.method == 'POST' and 'import' in request.POST:
            url = 'https://www.carwin.ru/japanauc/'
            response = requests.get(url)
            html_page = BeautifulSoup(response.text, 'lxml')
            urls_list = html_page.find('ul', 'pagination')
            urls_list = urls_list.find_all('a')

            for i in range(1, len(urls_list) + 1):

                response = requests.get(url + str(i))
                html_page = BeautifulSoup(response.text, 'lxml')
                links = html_page.find_all('a', 'pic')
                for link in links:
                    linked_list.append(link['href'])

            url = 'https://www.carwin.ru'
            print(url + linked_list[0])
            k = 0
            for link in linked_list:
                k += 1
                print(k)
                self.link_obr(url + link)
            messages.success(request, "База данных обновлена")

        return render(request, 'parser.html', {'response': 'success'})

    def link_obr(self, url):

        response = requests.get(url)
        html_page = BeautifulSoup(response.text, "lxml")
        number_of_auc = html_page.find('div', 'row_desc_middle')

        car = Car.objects.filter(auc_number=number_of_auc.text)

        if html_page.find('div', 'page_title') is not None and not car:

            Obj = Car_data(url)
            print(url)
            Obj.print()

            Obj.save_me_to_bd()
            del Obj
        else:

            print('Было, знаем!', url)
            # return


class Parser(object):
    html_page = None

    def __init__(self, url):
        response = requests.get(url)
        self.html_page = BeautifulSoup(response.text, "lxml")
        print(response)

    def parse_title(self):
        name = self.html_page.find('div', 'page_title').text
        # Достаем данные из excel
        file_path = 'cars_price.xlsx'
        workbook = load_workbook(filename=file_path)

        models_list = workbook.worksheets[1]
        mark_list = workbook.worksheets[0]

        models = []
        skip_first_row = False
        for row in models_list.iter_rows(values_only=True, min_row=2 if skip_first_row else 1):
            model = {
                'id': row[0],
                'mark': row[1],
                'model': row[2],
                'price': row[3],
            }
            models.append(model)
        marks = []
        for row in mark_list.iter_rows(values_only=True, min_row=2 if skip_first_row else 1):
            mark = {
                'id': row[0],
                'mark': row[1]
            }
            marks.append(mark)

        # Достаем данные из excel

        # Узнаем цену машины

        car_title = name.split()
        is_car = False
        price = 0
        for car_for_test in models:

            # if str(marks[int(car_for_test['mark'])-1]['mark']) == 'Acura':
            #     print(car_for_test, marks[int(car_for_test['mark'])-1]['mark'], 'true', int(car_for_test['mark']))
            # print(marks[int(car_for_test['mark'])-1]['mark'], car_title[0].lower())
            if (str(car_for_test['model']).lower() == car_title[1].lower()
                    and
                    str(marks[int(car_for_test['mark']) - 1]['mark']).lower() == car_title[0].lower()):
                is_car = True
                price = car_for_test['price']
        if is_car is False:
            price = 'Нет информации о цене'

        print(name, price, '______________________________________________________')
        price_name = {'name': name, 'price': price}
        return price_name

    def parse_auction_data(self):
        auction_data = self.html_page.find('div', 'col_left').text
        auction_data = [value for value in auction_data.split('\n') if value != '']
        return auction_data

    def parse_car_options(self):
        car_options = self.html_page.find('div', 'car_description')
        car_options = car_options.find_all('div', 'car_option')

        data_set = {'year_car': '', 'mileage': '', 'color': '', 'options': '', 'the_body': '', 'volume': '', 'cpp': '',
                    'estimation': ''}

        # Условие не трогать, работает и слава богу
        for el in range(0, len(car_options)):
            split_data_of_car = car_options[el].text.split()

            match split_data_of_car[0]:
                case 'Год':
                    if len(split_data_of_car) > 1:
                        data_set['year_car'] = split_data_of_car[1]
                case 'Пробег':
                    if len(split_data_of_car) > 1:
                        data_set['mileage'] = split_data_of_car[1]
                case 'Цвет':
                    if len(split_data_of_car) > 1:
                        data_set['color'] = split_data_of_car[1]
                case 'Опции':
                    if len(split_data_of_car) > 1:
                        data_set['options'] = split_data_of_car[1]
                case 'Кузов':
                    if len(split_data_of_car) > 1:
                        data_set['the_body'] = split_data_of_car[1]
                case 'Объем':
                    if len(split_data_of_car) > 1:
                        data_set['volume'] = split_data_of_car[1]
                case 'КПП':
                    if len(split_data_of_car) > 1:
                        data_set['cpp'] = split_data_of_car[1]
                case 'Оценка':
                    if len(split_data_of_car) > 1:
                        data_set['estimation'] = split_data_of_car[1]

        car_options = data_set
        return car_options

    def parse_content(self):
        content = self.html_page.find('div', 'content')
        content = content.find_all('td')

        data_set = {'cooling': '',
                    'set': '',
                    'result': '',
                    'start_price': '',
                    'transmission': '',
                    'location_auction': '',
                    'year': '',
                    'alt_color': '',
                    'condition': '',
                    'fuel': '',
                    'equipment': '',
                    'deadline_for_the_price_offer': '',
                    'day_of_the_event': '',
                    'number_of_sessions': ''}
        # Почему здесь по-другому, не знаю, но тоже работает и слава богу
        for el in range(0, len(content), 2):
            match content[el].text:
                case ' охлаждение ':
                    data_set['cooling'] = content[el + 1].text
                case ' комплектация ':
                    data_set['set'] = content[el + 1].text
                case ' результат ':
                    data_set['result'] = content[el + 1].text
                case ' старт ':
                    data_set['start_price'] = content[el + 1].text
                case ' коробка передач ':
                    data_set['transmission'] = content[el + 1].text
                case ' место проведения ':
                    data_set['location_auction'] = content[el + 1].text
                case ' год ':
                    data_set['year'] = content[el + 1].text
                case ' цвет ':
                    data_set['alt_color'] = content[el + 1].text
                case ' состояние ':
                    data_set['condition'] = content[el + 1].text
                case ' топливо ':
                    data_set['fuel'] = content[el + 1].text
                case ' оборудование ':
                    data_set['equipment'] = content[el + 1].text
                case ' конечный срок предложения цены ':
                    data_set['deadline_for_the_price_offer'] = content[el + 1].text
                case ' день проведения ':
                    data_set['day_of_the_event'] = content[el + 1].text
                case ' количество проведений ':
                    data_set['number_of_sessions'] = content[el + 1].text

        content = data_set
        return content

    def parse_image(self):
        image = self.html_page.find('div', 'my-gallery')
        image = image.find_all('img')
        form_data = list()

        for el in range(len(image)):
            form_data.append(image[el]['src'])
        image = form_data
        return image

    def parse_auc_list(self):
        auc_list = self.html_page.find('div', 'scheme_block')
        auc_list = auc_list.find('img')
        auc_list = auc_list['src']
        return auc_list


class Car_data(object):
    title = ''
    auction_data = ''
    car_options = ''
    content = ''
    auc_link = ''
    image = ''
    # auction_data
    auc_name = ''
    auc_number = ''
    auc_date = ''
    # car_options
    year_car = ''
    mileage = ''
    color = ''
    options = ''
    the_body = ''
    volume = ''
    cpp = ''
    estimation = ''
    # content
    cooling = ''
    set = ''
    result = ''
    start_price = ''
    transmission = ''
    location_auction = ''
    year = ''
    alt_color = ''
    condition = ''
    fuel = ''
    equipment = ''
    deadline_for_the_price_offer = ''
    day_of_the_event = ''
    number_of_sessions = ''

    auc_list = ''

    price = ''

    def __init__(self, url):
        parser = Parser(url)
        self.auc_link = url
        # для Car_of_page
        price_name = parser.parse_title()
        self.title = price_name['name']
        self.price = price_name['price']
        self.auction_data = parser.parse_auction_data()
        self.car_options = parser.parse_car_options()
        self.content = parser.parse_content()
        self.image = parser.parse_image()

        # для Car_data

        # auction_data
        self.auc_name = parser.parse_auction_data()[0]
        self.auc_number = parser.parse_auction_data()[1]
        self.auc_date = parser.parse_auction_data()[2]

        # car_options
        form_data = parser.parse_car_options()

        self.year_car = form_data['year_car']
        self.mileage = form_data['mileage']
        self.color = form_data['color']
        self.options = form_data['options']
        self.the_body = form_data['the_body']
        self.volume = form_data['volume']
        self.cpp = form_data['cpp']
        self.estimation = form_data['estimation']

        # content
        form_data = parser.parse_content()

        self.cooling = form_data['cooling']
        self.set = form_data['set']
        self.result = form_data['result']
        self.start_price = form_data['start_price']
        self.transmission = form_data['transmission']
        self.location_auction = form_data['location_auction']
        self.year = form_data['year']
        self.alt_color = form_data['alt_color']
        self.condition = form_data['condition']
        self.fuel = form_data['fuel']
        self.equipment = form_data['equipment']
        self.deadline_for_the_price_offer = form_data['deadline_for_the_price_offer']
        self.day_of_the_event = form_data['day_of_the_event']
        self.number_of_sessions = form_data['number_of_sessions']

        self.auc_list = parser.parse_auc_list()

    def print(self):
        print('название машины', self.title, 'аукцион', self.auction_data, 'основное про машину', self.car_options,
              'таблица', self.content, sep='\n', end='\n')
        print('картинки', self.image, end='\n')
        print(self.auc_name, self.auc_number, self.auc_date, sep='\n', end='\n')
        print(self.year_car, self.mileage, self.color, self.options, self.the_body, self.volume, self.cpp,
              self.estimation, sep='\n', end='\n')
        print(self.cooling, self.condition, self.fuel, self.equipment)

    def __del__(self):
        print('Удален')

    def save_me_to_bd(self):
        new_car_new = Car.objects.create(
            auc_link=self.auc_link,
            title=self.title,
            auc_name=self.auc_name,
            auc_number=self.auc_number,
            auc_date=self.auc_date,
            year_car=self.year_car,
            mileage=self.mileage,
            color=self.color,
            options=self.options,
            the_body=self.the_body,
            volume=self.volume,
            cpp=self.cpp,
            estimation=self.estimation,
            cooling=self.cooling,
            set=self.set,
            result=self.result,
            start_price=self.start_price,
            transmission=self.transmission,
            location_auction=self.location_auction,
            year=self.year,
            alt_color=self.alt_color,
            condition=self.condition,
            fuel=self.fuel,
            equipment=self.equipment,
            deadline_for_the_price_offer=self.deadline_for_the_price_offer,
            day_of_the_event=self.day_of_the_event,
            number_of_sessions=self.number_of_sessions,
            auc_list=self.auc_list,
            price=self.price
        )
        for el in range(len(self.image)):
            PhotoCar.objects.create(id_car=new_car_new, photo=self.image[el])

        print(new_car_new)


class BuhgalterPageView(TemplateView):
    template_name = "buhgalter/buhgalter.html"

    def get(self, request, *args, **kwargs):
        invoices = Invoice.objects.all()
        print('Тут должен быть инвойс')
        print(invoices)
        return render(request, 'buhgalter/buhgalter.html', {'invoices': invoices})


class BuhgalterInvoicePageView(TemplateView):
    template_name = "buhgalter/invoice.html"

    def get(self, request, *args, **kwargs):
        invoice = Invoice.objects.get(id_invoice=kwargs.get('invoice_id'))
        invoice_data = Invoice.objects.all()
        form = InvoiceForm()
        form.fields['id_invoice'].widget.attrs.update({'value': invoice.id_invoice})
        form.fields['payer'].widget.attrs.update({'value': invoice.payer})
        form.fields['seller'].widget.attrs.update({'value': invoice.seller})
        form.fields['date_form'].widget.attrs.update({'value': invoice.date_form})
        form.fields['date_pay'].widget.attrs.update({'value': invoice.date_pay})
        form.fields['sum'].widget.attrs.update({'value': invoice.sum})
        form.fields['check_document'].widget.attrs.update({'value': invoice.check_document})
        form.fields['assigning'].widget.attrs.update({'value': invoice.assigning})
        form.fields['scan'].widget.attrs.update({'value': invoice.scan})
        form.fields['type'].widget.attrs.update({'value': invoice.type})

        return render(request, 'buhgalter/invoice.html',
                      {'invoice': invoice, 'invoice_data': invoice_data, 'form': form})

    def post(self, request, *args, **kwargs):
        invoices = Invoice.objects.all()
        if request.method == 'POST':
            form = InvoiceForm(request.POST)
            print('Валидная или инвалидная форма', form.is_valid())
            print(form.errors)
            if form.is_valid():
                form.update()
                messages.success(request, "Данные обновлены")
            else:
                messages.error(request, "Некорректная форма")
        else:
            form = InvoiceForm()
        return render(request, 'buhgalter/buhgalter.html', {'invoices': invoices})


class BuhgalterNewInvoicePageView(TemplateView):
    template_name = "buhgalter/new_invoice.html"

    def get(self, request, *args, **kwargs):
        if request.method == 'GET':
            form = InvoiceForm()
            return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        invoices = Invoice.objects.all()
        if request.method == 'POST':
            print('Запрос пришел')
            form = NewInvoiceForm(request.POST)
            print(form.is_valid())
            print(form.errors)
            if form.is_valid():
                form.save()
                messages.success(request, "Счет на оплату сохранен")
                return render(request, 'buhgalter/new_invoice.html', {'form': form})
            else:
                messages.error(request, "Некорректная форма")
            # form.save()
            # messages.info(request, "Счет на оплату сохранен")
        else:
            form = NewInvoiceForm()
        return render(request, 'buhgalter/buhgalter.html', {'invoices': invoices})


def orders(request):
    user_id = request.user.id
    user = Worker.objects.get(pk=user_id)
    user_passport = user.passport.split()
    if request.method == 'GET' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        # Если это AJAX-запрос, обрабатываем его
        status = request.GET.get('status')
        full_name = request.GET.get('last_name')
        orders = Order.objects.all()
        start_date_str = request.GET.get('start_date')
        end_date_str = request.GET.get('end_date')

        if status:
            # Фильтрация по статусу
            orders = orders.filter(order_status=status)

        if full_name:
            # Фильтрация по фамилии, имени или отчеству
            parts = full_name.split()
            condition = Q()
            for part in parts:
                condition |= Q(id_customer__last_name_client__icontains=part) | \
                             Q(id_customer__first_name_client__icontains=part) | \
                             Q(id_customer__patronymic_client__icontains=part)
            orders = orders.filter(condition)

        if start_date_str and end_date_str:
            # Фильтрация заказов по дате
            orders = orders.filter(date_start__range=[start_date_str, end_date_str])


        if user.job_title == 'Клиент':
            orders = Order.objects.filter(id_customer__passport_number=user_passport[1], id_customer__passport_series=user_passport[0])
        context = {
            'orders': orders
        }
        html = render_to_string('order_table.html', context)
        return JsonResponse({'html': html})

    # Если это не AJAX-запрос, возвращаем страницу заказов целиком
    orders = Order.objects.all()
    if user.job_title == 'Клиент':
        orders = Order.objects.filter(id_customer__passport_number=user_passport[1],
                                      id_customer__passport_series=user_passport[0])
    context = {
        'orders': orders
    }
    return render(request, 'orders.html', context)


def customers(request):
    if request.method == 'GET' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        # Если это AJAX-запрос, обрабатываем его
        phone = request.GET.get('phone')
        full_name = request.GET.get('last_name')
        customers = Customer.objects.all()
        passport = request.GET.get('passport')

        print(phone, full_name, passport)

        if phone:
            # Фильтрация по статусу
            customers = customers.filter(telephone__icontains=phone)


        if full_name:
            # Фильтрация по фамилии, имени или отчеству
            parts = full_name.split()
            condition = Q()
            for part in parts:
                condition |= Q(last_name_client__icontains=part) | \
                             Q(first_name_client__icontains=part) | \
                             Q(patronymic_client__icontains=part)
            customers = customers.filter(condition)

        if passport:
            parts = passport.split()
            condition = Q()
            for part in parts:
                condition &= Q(passport_series__icontains=part) | Q(passport_number__icontains=part)
            customers = customers.filter(condition)

        context = {
            'customers': customers
        }
        print(context)
        html = render_to_string('customer_table.html', context)
        return JsonResponse({'html': html})

        # Если это не AJAX-запрос, возвращаем страницу заказов целиком
    customers = Customer.objects.all()
    context = {
        'customers': customers
    }
    return render(request, 'customers.html', context)



def transport_companies(request):
    if request.method == 'GET' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        name = request.GET.get('name', '')
        transport_companies = TransportCompany.objects.all()

        if name:
            transport_companies = transport_companies.filter(title__icontains=name)

        context = {
            'transport_companies': transport_companies
        }
        html = render_to_string('transport_companies_table.html', context)
        return JsonResponse({'html': html})

    # Загружаем начальные данные при первой загрузке страницы
    transport_companies = TransportCompany.objects.all()
    print(transport_companies)
    context = {
        'transport_companies': transport_companies
    }
    return render(request, 'transport_companies.html', context)

def catalog(request):
    if request.method == 'GET' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        carName = request.GET.get('car_name')
        engine_volume = request.GET.get('engine_volume')
        year = request.GET.get('year')
        ratings = request.GET.getlist('rating')
        mileage = request.GET.get('mileage')
        car = Car.objects.all()
        print('/', carName, '/', engine_volume, year, ratings, mileage)

        if carName:
            carName = carName.upper()
            print('/',carName, '/', engine_volume, year, ratings, mileage)
            car = car.filter(title__icontains=carName)

        if engine_volume:
            car = car.filter(volume=engine_volume)

        if year:
            car = car.filter(year_car=year)

        if ratings:
            car = car.filter(estimation__in=ratings)

        if mileage:
            car = car.filter(mileage=mileage)

        for el in car:
            el.image = PhotoCar.objects.filter(id_car=el.id_car)[:1][0].photo



        context = {
            'cars': car,
        }
        html = render_to_string('catalog_table.html', context)
        return JsonResponse({'html': html})

    cars = Car.objects.all()
    for car in cars:
        car.image = PhotoCar.objects.filter(id_car=car.id_car)[:1][0].photo
    context = {
        'cars': cars
    }
    return render(request, 'catalog.html', context)

def duties(request):
    if request.method == 'GET' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        start_date = request.GET.get('start_date')

        all_dates = Duty.objects.values_list('date_of_action', flat=True).distinct()

        # Преобразуем QuerySet в список дат
        unique_dates_list = list(all_dates)

        # Убираем None из списка, если есть
        unique_dates_list = [date for date in unique_dates_list if date is not None]

        # Сортируем даты по возрастанию
        unique_dates_list.sort()

        duties = Duty.objects.all()
        if start_date:
            # Найти последнюю дату пошлины, которая меньше или равна указанной дате
            latest_date = Duty.objects.filter(date_of_action__lte=start_date).aggregate(Max('date_of_action'))['date_of_action__max']
            if latest_date:
                duties = duties.filter(date_of_action=latest_date)
        print(duties)
        context = {
            'duties': duties
        }
        html = render_to_string('duties_table.html', context)
        return JsonResponse({'html': html})

    duties = Duty.objects.all()
    context = {
        'duties': duties
    }
    return render(request, 'duties.html', context)


def customs_dutys(request):
    if request.method == 'GET' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        start_date = request.GET.get('start_date')
        typeof = request.GET.get('status')
        print(typeof)
        all_dates = CustomsDuty.objects.values_list('date_of_action', flat=True).distinct()

        # Преобразуем QuerySet в список дат
        unique_dates_list = list(all_dates)

        # Убираем None из списка, если есть
        unique_dates_list = [date for date in unique_dates_list if date is not None]

        # Сортируем даты по возрастанию
        unique_dates_list.sort()

        customs_dutys = CustomsDuty.objects.all()
        if start_date:
            # Найти последнюю дату пошлины, которая меньше или равна указанной дате
            latest_date = CustomsDuty.objects.filter(date_of_action__lte=start_date).aggregate(Max('date_of_action'))['date_of_action__max']
            if latest_date:
                customs_dutys = customs_dutys.filter(date_of_action=latest_date)
            else:
                customs_dutys = CustomsDuty.objects.none()  # Если нет подходящей даты, возвращаем пустой QuerySet

        if typeof:
            customs_dutys = customs_dutys.filter(type=typeof)
        context = {
            'customs_dutys': customs_dutys
        }
        html = render_to_string('customs_dutys_table.html', context)
        return JsonResponse({'html': html})

    customs_dutys = CustomsDuty.objects.all()
    context = {
        'customs_dutys': customs_dutys
    }
    return render(request, 'customs_dutys.html', context)


def excises(request):
    if request.method == 'GET' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        start_date = request.GET.get('start_date')

        all_dates = Excise.objects.values_list('date_of_action', flat=True).distinct()

        # Преобразуем QuerySet в список дат
        unique_dates_list = list(all_dates)

        # Убираем None из списка, если есть
        unique_dates_list = [date for date in unique_dates_list if date is not None]

        # Сортируем даты по возрастанию
        unique_dates_list.sort()

        excises = Excise.objects.all()
        if start_date:
            # Найти последнюю дату акциза, которая меньше или равна указанной дате
            latest_date = Excise.objects.filter(date_of_action__lte=start_date).aggregate(Max('date_of_action'))['date_of_action__max']
            if latest_date:
                excises = excises.filter(date_of_action=latest_date)
            else:
                excises = Excise.objects.none()  # Если нет подходящей даты, возвращаем пустой QuerySet

        context = {
            'excises': excises
        }
        print(excises)
        html = render_to_string('excises_table.html', context)
        return JsonResponse({'html': html})

    excises = Excise.objects.all()
    context = {
        'excises': excises
    }
    return render(request, 'excises.html', context)


def download_all_documents(request, order_id):

    order = Order.objects.get(id_order=order_id)
    path_def_ved = OrderInOrdersPageView().create_dev_ved(order.pk)['title']

    # document = Document(doc_path)
    documents = [
        order.sbts,
        order.ptd,
        order.contract,
        order.defective_statement,
        order.export_certificate,
    ]
    if not order.contract:
        path_contract = OrderInOrdersPageView().create_contract(order.pk)['title']
        documents.append(path_contract)
    if not order.defective_statement:
        path_def_ved = OrderInOrdersPageView().create_dev_ved(order.pk)['title']
        documents.append(path_def_ved)




    # documents.extend(new_documents)
    # Create a zip file in memory
    zip_subdir = f"order_{order_id}_documents"
    zip_filename = f"{zip_subdir}.zip"

    s = HttpResponse(content_type="application/zip")
    s['Content-Disposition'] = f'attachment; filename={zip_filename}'

    with zipfile.ZipFile(s, 'w') as zf:
        for doc in documents:
            if doc:
                if isinstance(doc, str):
                    arcname = os.path.join(zip_subdir, os.path.basename(doc))
                    zf.write(doc, arcname)
                else:
                    doc_path = os.path.join(os.getcwd(), doc.path)
                    arcname = os.path.join(zip_subdir, os.path.basename(doc_path))
                    zf.write(doc_path, arcname)

    return s


def order_photo(request, order_id):
    print(order_id)  # Логирование для проверки
    order = get_object_or_404(Order, pk=order_id)
    if request.method == 'POST':
        if 'photos' in request.FILES:
            photo = PhotoGallery.objects.create(order=order, photo=request.FILES.get('photos'))
            photo.save()
            order = get_object_or_404(Order, pk=order_id)
            context = {
                'order': order,
            }
            html = render_to_string('photo_gallery.html', context)
            return JsonResponse({'html': html})
        else:
            return JsonResponse({'error': 'Invalid form'}, status=400)
    return JsonResponse({'error': 'Invalid request'}, status=400)


class OrderTransportView(TemplateView):
    template_name = 'order_transport.html'

    def get(self, request, *args, **kwargs):
        order = Order.objects.all()
        cars = Car.objects.all()
        transport_companies = TransportCompany.objects.all()
        return render(request, 'order_transport.html', {'cars': cars, 'transport_companies': transport_companies})

    def post(self, request, *args, **kwargs):
        if request.method == 'POST':
            print(request.POST)
        return django.http.HttpResponseRedirect(reverse('order_transport'))


def update_transport_companies(request):
    selected_cars = request.GET.getlist('cars[]')
    print(selected_cars)
    # Здесь можно добавить логику для фильтрации компаний в зависимости от выбранных машин
    # Для примера просто возвращаем все компании
    transport_companies = TransportCompany.objects.filter(pk=1)
    context = {
        'transport_companies': transport_companies,
    }
    html = render_to_string('transport_select.html', context)
    return JsonResponse({'html': html})